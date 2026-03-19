from __future__ import annotations

import io
import os
from collections.abc import AsyncGenerator

import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import create_async_engine, async_sessionmaker
from sqlmodel import SQLModel
from sqlmodel.ext.asyncio.session import AsyncSession

# ---------------------------------------------------------------------------
# Test database configuration
# ---------------------------------------------------------------------------

TEST_DATABASE_URL = "sqlite+aiosqlite:///./test.db"

# Override environment variables BEFORE importing any application modules so
# that ``get_settings()`` (which is cached with ``@lru_cache``) picks up the
# test values the first time it is called.
os.environ["DATABASE_URL"] = TEST_DATABASE_URL
os.environ["SECRET_KEY"] = "test-secret-key-at-least-32-characters-long!"
os.environ["REDIS_URL"] = "redis://localhost:6379"
os.environ["R2_ACCOUNT_ID"] = "test-account-id"
os.environ["R2_ACCESS_KEY_ID"] = "test-access-key"
os.environ["R2_SECRET_ACCESS_KEY"] = "test-secret-key"
os.environ["R2_BUCKET_NAME"] = "test-bucket"
os.environ["ADZUNA_APP_ID"] = "test-adzuna-id"
os.environ["ADZUNA_APP_KEY"] = "test-adzuna-key"
os.environ["USAJOBS_API_KEY"] = "test-usajobs-key"
os.environ["USAJOBS_EMAIL"] = "test@example.com"
os.environ["APIFY_API_KEY"] = "test-apify-key"

# Now it is safe to import the app and its dependencies.
from app.config import get_settings  # noqa: E402
from app.database import get_db  # noqa: E402
from app.main import app  # noqa: E402

# ---------------------------------------------------------------------------
# Async engine and session factory scoped to the test module
# ---------------------------------------------------------------------------

_test_engine = create_async_engine(
    TEST_DATABASE_URL,
    echo=False,
    connect_args={"check_same_thread": False},
)

_test_session_factory = async_sessionmaker(
    bind=_test_engine,
    class_=AsyncSession,
    expire_on_commit=False,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest_asyncio.fixture
async def session() -> AsyncGenerator[AsyncSession, None]:
    """Provide a clean async database session for each test.

    Creates all tables before yielding, and drops them after the test to
    ensure full isolation between tests.
    """
    async with _test_engine.begin() as conn:
        await conn.run_sync(SQLModel.metadata.create_all)

    async with _test_session_factory() as sess:
        try:
            yield sess
            await sess.commit()
        except Exception:
            await sess.rollback()
            raise

    async with _test_engine.begin() as conn:
        await conn.run_sync(SQLModel.metadata.drop_all)


@pytest_asyncio.fixture
async def client(session: AsyncSession) -> AsyncGenerator[AsyncClient, None]:
    """Provide an async HTTP test client with overridden DB dependency.

    The database dependency is replaced so that every request within a test
    uses the same ``session`` (and therefore the same in-memory state).
    """

    async def _override_get_db() -> AsyncGenerator[AsyncSession, None]:
        yield session

    app.dependency_overrides[get_db] = _override_get_db

    transport = ASGITransport(app=app)  # type: ignore[arg-type]
    async with AsyncClient(transport=transport, base_url="http://testserver") as ac:
        yield ac

    app.dependency_overrides.clear()


@pytest_asyncio.fixture
async def auth_headers(client: AsyncClient) -> dict[str, str]:
    """Register a test user and return Authorization headers with a valid JWT.

    The user is created via the /register endpoint so that all service-layer
    side-effects (profile creation, etc.) are exercised.
    """
    response = await client.post(
        "/api/v1/auth/register",
        json={
            "email": "testuser@example.com",
            "password": "securepassword123",
            "full_name": "Test User",
        },
    )
    assert response.status_code == 201, f"Registration failed: {response.text}"
    data = response.json()
    access_token = data["access_token"]
    return {"Authorization": f"Bearer {access_token}"}


@pytest_asyncio.fixture
async def test_jobs(session: AsyncSession) -> list:
    """Create test jobs in the database for job-related tests.

    Returns a list of 3 Job model instances with varying attributes
    to exercise filtering by source, job_type, is_remote, etc.
    """
    from datetime import UTC, datetime

    from app.models.job import Job

    jobs: list[Job] = []
    for i in range(3):
        job = Job(
            external_id=f"test-job-{i}",
            source="remoteok" if i < 2 else "themuse",
            title=f"Test Job {i}" if i < 2 else "Data Analyst",
            company=f"Company {i}",
            location="Remote" if i == 0 else "New York" if i == 1 else None,
            is_remote=i == 0,
            salary_min=80000 + i * 20000,
            salary_max=120000 + i * 20000,
            salary_currency="USD",
            description=f"Description for job {i}",
            job_type="full-time" if i < 2 else "contract",
            url=f"https://example.com/job/{i}",
            tags=["python", "react"] if i < 2 else ["sql", "excel"],
            raw_data={"test": True},
            posted_at=datetime.now(UTC),
            is_active=True,
        )
        session.add(job)
        jobs.append(job)
    await session.flush()
    return jobs


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Create minimal PDF bytes with resume-like content."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "John Doe\njohn@example.com\n555-123-4567\n\nSkills\nPython, TypeScript, React",
    )
    content = doc.tobytes()
    doc.close()
    return content


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Create minimal DOCX bytes with resume-like content."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane@example.com")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, React, SQL")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
